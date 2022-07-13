from typing import BinaryIO

from django.shortcuts import render,redirect
from django.http import HttpResponse
from .forms import DocumentForm
from .models import Mydocument
import docx as dc
from django.core.files import File
from pathlib import Path
from django.core.files.storage import FileSystemStorage
# Create your views here.
def SS_Merge(response, n, f1,f2,f3,m):
    import os
    url_str=r"C:\Users\Dell\PycharmProjects\WordProcessor_web"
    loc = url_str+ m
    print("location of merged file",loc)
    loc_1 = url_str+f1
    loc_2 = url_str+f2
    loc_3 = url_str+f3

    doc = dc.Document(loc)
    doc1 = dc.Document(loc_1)
    doc2 = dc.Document(loc_2)
    doc3 = dc.Document(loc_2)
    List_ss1_col1 = []
    List_ss1_col2 = []
    List_ss2_col1 = []
    List_ss2_col2 = []
    List_ss3_col1 = []
    List_ss3_col2 = []
    # Lists to keep each files two column's values.

    for table1 in doc1.tables:
        for index in range(2, 20, 1):
            if (index == 8 or index == 13):
                List_ss1_col1.append('')
                List_ss1_col2.append('')
                continue

            # reading values from each cells for each of the six rows and appending to two column lists
            cell_id = table1.rows[index].cells
            List_ss1_col1.append(cell_id[1].text)
            List_ss1_col2.append(cell_id[2].text)
    print(List_ss1_col1)
    print(List_ss1_col2)

    for table2 in doc2.tables:
        for index in range(2, 20, 1):
            if (index == 8 or index == 13):
                List_ss2_col1.append('')
                List_ss2_col2.append('')
                continue
            # reading values for separate cells of second file and appending two column's data to two separate lists.
            cell_id = table2.rows[index].cells
            List_ss2_col1.append(cell_id[1].text)
            List_ss2_col2.append(cell_id[2].text)
    print(List_ss2_col1)
    print(List_ss2_col2)

    for table3 in doc3.tables:
        for index in range(2, 20, 1):
            if (index == 8 or index == 13):
                List_ss3_col1.append('')
                List_ss3_col2.append('')
                continue
            # reading and appending data to the list for 3rd file.
            cell_id = table3.rows[index].cells
            List_ss3_col1.append(cell_id[1].text)
            List_ss3_col2.append(cell_id[2].text)
    print(List_ss3_col1)
    print(List_ss3_col2)

    for table in doc.tables:
        for i in range(6):
            # Reading data from respective lists and writing them on specified locations in the merged file.
            column = table.rows[2 + i * 3].cells
            # data appending for the first row (row 2 for i=0)
            # data writing to the respective cells for each of the six position of the column list of each file
            column[1].text = List_ss1_col1[i]
            # data writing to the first column from first position of the list
            column[2].text = List_ss1_col2[i]
            # incrementing rows index by three each after a iteration of putting values from the each locations of list.
            column = table.rows[3 + i * 3].cells
            column[1].text = List_ss2_col1[i]
            column[2].text = List_ss2_col2[i]
            column = table.rows[4 + i * 3].cells
            column[1].text = List_ss3_col1[i]
            column[2].text = List_ss3_col2[i]
        for i in range(4):
            column = table.rows[21 + i * 3].cells
            column[1].text = List_ss1_col1[i + 7]
            column[2].text = List_ss1_col2[i + 7]
            column = table.rows[22 + i * 3].cells
            column[1].text = List_ss2_col1[i + 7]
            column[2].text = List_ss2_col2[i + 7]
            column = table.rows[23 + i * 3].cells
            column[1].text = List_ss3_col1[i + 7]
            column[2].text = List_ss3_col2[i + 7]
        for i in range(6):
            column = table.rows[34 + i * 3].cells
            column[1].text = List_ss1_col1[i + 12]
            column[2].text = List_ss1_col2[i + 12]
            column = table.rows[35 + i * 3].cells
            column[1].text = List_ss2_col1[i + 12]
            column[2].text = List_ss2_col2[i + 12]
            column = table.rows[36 + i * 3].cells
            column[1].text = List_ss3_col1[i + 12]
            column[2].text = List_ss3_col2[i + 12]
    file_name = r'C:\Users\Dell\PycharmProjects\WordProcessor_web\media\ssmerged.docx'
    doc.save(file_name)
    print(loc)
    file = Mydocument.objects.get(id=n)
    f_name = r'C:/Users/Dell/PycharmProjects/WordProcessor_web/media/ssmerged.docx'
    path = Path(f_name)
    with path.open(mode = 'rb') as f:
        file_opened= File(f, name = path.name)
        file.file_processed.specs = file_opened
        file.save()
        Mydocument.objects.filter(id =n).update(file_processed=file_opened)

        print("This file is ssmerged file:",file.file_processed.url)


    #Trying to update the file at database with the new one"



def home(response):
    con = {}
    if response.method == 'POST':
        form = DocumentForm(response.POST, response.FILES)
        if form.is_valid():
            new_doc= form.save()
            file1_name = new_doc.file1.url
            file2_name = new_doc.file2.url
            file3_name = new_doc.file3.url
            merged_url = new_doc.file_processed.url

            number =new_doc.pk
            print(number)
            print("url",file2_name)
            print("url",file1_name)
            SS_Merge(response,number,file1_name,file2_name,file3_name,merged_url)
            #return redirect('download')
            #print("Sucessful File Merged To",merged_url)

            doc1 = Mydocument.objects.get(id = number)
            con['url'] = doc1.file_processed.url
            print("File merged url is :",doc1.file_processed.url)
            return render(response,'download.html', con)

    else:
        form = DocumentForm()
        return render(response,'upload.html',{'form':form})
def download( response):
    pass